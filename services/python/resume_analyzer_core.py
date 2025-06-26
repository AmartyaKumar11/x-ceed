"""
X-ceed Resume Analyzer - Core RAG Engine
Modular chatbot that can be integrated into any project
"""

import os
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.schema import Document
import tempfile
from typing import List, Dict, Optional, Any

class ResumeAnalyzerCore:
    """
    Core RAG engine for resume analysis that can be integrated into any project.
    Handles document processing, vector storage, and AI conversations.
    """
    
    def __init__(self, groq_api_key: str = None):
        """
        Initialize the core analyzer
        
        Args:
            groq_api_key: Groq API key (if not provided, will load from environment)
        """
        # Load environment variables
        load_dotenv()
        
        self.groq_api_key = groq_api_key or os.getenv("GROQ_API_KEY")
        if not self.groq_api_key:
            raise ValueError("Please provide Groq API key or set GROQ_API_KEY environment variable")
        
        # Initialize the LLM
        self.llm = ChatGroq(
            temperature=0.1,
            groq_api_key=self.groq_api_key,
            model_name="llama-3.1-8b-instant"
        )
        
        # Initialize embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        self.vectorstore = None
        self.conversation_chain = None
        self.chat_history = []
        
    def process_documents(self, resume_text: str, job_description_text: str) -> bool:
        """
        Process resume and job description texts to create vector store
        
        Args:
            resume_text: Full text content of the resume
            job_description_text: Full text content of the job description
            
        Returns:
            bool: True if processing successful, False otherwise
        """
        try:
            documents = []
            
            if resume_text:
                resume_chunks = self.text_splitter.split_text(resume_text)
                for chunk in resume_chunks:
                    documents.append(Document(
                        page_content=chunk,
                        metadata={"source": "resume", "type": "resume_content"}
                    ))
            
            if job_description_text:
                job_chunks = self.text_splitter.split_text(job_description_text)
                for chunk in job_chunks:
                    documents.append(Document(
                        page_content=chunk,
                        metadata={"source": "job_description", "type": "job_requirements"}
                    ))
            
            if documents:
                # Create a temporary directory for ChromaDB
                persist_directory = tempfile.mkdtemp()
                self.vectorstore = Chroma.from_documents(
                    documents=documents, 
                    embedding=self.embeddings,
                    persist_directory=persist_directory
                )
                
                # Create conversation chain
                self._create_conversation_chain()
                return True
            return False
            
        except Exception as e:
            print(f"Error processing documents: {str(e)}")
            return False
    
    def _create_conversation_chain(self):
        """Create conversation chain with memory"""
        if self.vectorstore:
            memory = ConversationBufferMemory(
                memory_key='chat_history',
                return_messages=True,
                output_key='answer'
            )
            
            self.conversation_chain = ConversationalRetrievalChain.from_llm(
                llm=self.llm,
                retriever=self.vectorstore.as_retriever(search_kwargs={"k": 3}),
                memory=memory,
                return_source_documents=True,
                verbose=False
            )
    
    def ask_question(self, question: str) -> Dict[str, Any]:
        """
        Ask a question about the resume and job description
        
        Args:
            question: The question to ask
            
        Returns:
            Dict containing answer, sources, and metadata
        """
        if not self.conversation_chain:
            return {
                "answer": "Please process documents first using process_documents() method.",
                "sources": [],
                "error": "No documents processed"
            }
        
        try:
            response = self.conversation_chain({'question': question})
            
            # Store in chat history
            self.chat_history.append({
                "question": question,
                "answer": response['answer'],
                "sources": response.get('source_documents', [])
            })
            
            return {
                "answer": response['answer'],
                "sources": response.get('source_documents', []),
                "success": True
            }
            
        except Exception as e:
            return {
                "answer": f"Error processing question: {str(e)}",
                "sources": [],
                "error": str(e)
            }
    
    def get_quick_analysis(self, analysis_type: str = "match") -> Dict[str, Any]:
        """
        Get quick analysis results
        
        Args:
            analysis_type: Type of analysis ('match', 'skills', 'improvements')
            
        Returns:
            Dict containing analysis results
        """
        questions = {
            "match": "Provide a detailed analysis of how well my resume matches the job requirements. Give me a percentage match and explain the key alignments and gaps.",
            "skills": "What skills and qualifications mentioned in the job description are missing from my resume? Provide specific recommendations.",
            "improvements": "Give me 5 specific suggestions to improve my resume for this job, including keywords I should add and sections I should enhance."
        }
        
        if analysis_type not in questions:
            return {"error": "Invalid analysis type. Use 'match', 'skills', or 'improvements'"}
        
        return self.ask_question(questions[analysis_type])
    
    def get_comprehensive_analysis(self) -> Dict[str, Any]:
        """
        Get comprehensive analysis similar to the JavaScript version
        
        Returns:
            Dict containing comprehensive analysis results
        """
        try:
            analysis_prompt = """
            Provide a comprehensive analysis of how well my resume matches this job description. 
            Return the analysis in the following structured format:

            OVERALL MATCH SCORE: [0-100]%
            MATCH LEVEL: [Excellent/Good/Fair/Poor]
            SUMMARY: [Brief summary of match quality]

            KEY STRENGTHS:
            - [List 3-5 key strengths with evidence]

            MATCHING SKILLS:
            - [List skills found in both resume and job description]

            MISSING SKILLS:
            - [List critical skills mentioned in job description but missing from resume]

            EXPERIENCE ANALYSIS:
            - Relevant Experience: [Years and specific roles]
            - Experience Gaps: [What experience is missing]

            IMPROVEMENT SUGGESTIONS:
            - [List 5 specific suggestions for improvement]

            COMPETITIVE ADVANTAGES:
            - [What makes this candidate stand out]

            INTERVIEW PREPARATION:
            - Strengths to Highlight: [Key points to emphasize]
            - Areas to Address: [Potential weaknesses to prepare for]

            Provide specific, actionable insights with evidence from the resume.
            """
            
            response = self.ask_question(analysis_prompt)
            
            if response.get('success'):
                return {
                    "success": True,
                    "analysis": response['answer'],
                    "comprehensive": True,
                    "timestamp": "2025-06-15T18:00:00.000Z"
                }
            else:
                return {
                    "success": False,
                    "error": response.get('error', 'Analysis failed'),
                    "analysis": response.get('answer', 'No analysis available')
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "analysis": "Failed to generate comprehensive analysis"
            }
    
    def get_chat_history(self) -> List[Dict]:
        """Get the complete chat history"""
        return self.chat_history
    
    def clear_session(self):
        """Clear the current session"""
        self.vectorstore = None
        self.conversation_chain = None
        self.chat_history = []
    
    def is_ready(self) -> bool:
        """Check if the analyzer is ready to answer questions"""
        return self.conversation_chain is not None


# Utility functions for document processing
class DocumentProcessor:
    """Utility class for processing different document types"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_file_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        except Exception as e:
            return f"Error reading PDF: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(docx_file_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(docx_file_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(txt_file_bytes: bytes) -> str:
        """Extract text from TXT bytes"""
        try:
            return txt_file_bytes.decode('utf-8')
        except Exception as e:
            return f"Error reading TXT: {str(e)}"
    
    @classmethod
    def process_file(cls, file_bytes: bytes, file_type: str) -> str:
        """
        Process file based on type
        
        Args:
            file_bytes: File content as bytes
            file_type: File type ('pdf', 'docx', 'txt')
            
        Returns:
            Extracted text content
        """
        processors = {
            'pdf': cls.extract_text_from_pdf,
            'docx': cls.extract_text_from_docx,
            'txt': cls.extract_text_from_txt
        }
        
        processor = processors.get(file_type.lower())
        if processor:
            return processor(file_bytes)
        else:
            return f"Unsupported file type: {file_type}"


if __name__ == "__main__":
    # Example usage
    analyzer = ResumeAnalyzerCore()
    
    # Example texts (you would get these from your document processing)
    sample_resume = "John Doe, Software Engineer with 5 years experience in Python, React, and AWS..."
    sample_job = "We are looking for a Senior Software Engineer with Python, React, and cloud experience..."
    
    # Process documents
    success = analyzer.process_documents(sample_resume, sample_job)
    
    if success:
        # Ask questions
        response = analyzer.ask_question("How well does my resume match this job?")
        print(response["answer"])
        
        # Get quick analysis
        match_analysis = analyzer.get_quick_analysis("match")
        print(match_analysis["answer"])
