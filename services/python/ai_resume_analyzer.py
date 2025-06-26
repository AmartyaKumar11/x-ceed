"""
AI Resume Analysis and Candidate Shortlisting Service - Gemini Version
Analyzes resumes and ranks candidates for job positions using Google Gemini AI
"""

import os
import json
import re
import io
from typing import List, Dict, Any
from dataclasses import dataclass
import google.generativeai as genai
from datetime import datetime
from dotenv import load_dotenv

# Load environment variables from .env.local
load_dotenv('.env.local')

@dataclass
class CandidateScore:
    candidate_id: str
    candidate_name: str
    overall_score: float
    skill_match_score: float
    experience_score: float
    education_score: float
    criteria_analysis: Dict[str, Any]
    strengths: List[str]
    weaknesses: List[str]
    recommendation: str

class AIResumeAnalyzer:
    def __init__(self):
        # Set up Gemini
        gemini_api_key = os.getenv('GEMINI_API_KEY')
        if not gemini_api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        
        genai.configure(api_key=gemini_api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')
    
    def analyze_resume_for_job(self, resume_text: str, job_description: str, job_requirements: Dict) -> Dict[str, Any]:
        """
        Analyze a single resume against job requirements using Gemini AI
        """
        try:
            prompt = self._create_analysis_prompt(resume_text, job_description, job_requirements)
            
            response = self.model.generate_content(prompt)
            
            analysis_text = response.text
            return self._parse_ai_response(analysis_text)
            
        except Exception as e:
            print(f"Error analyzing resume: {e}")
            return self._create_fallback_analysis()
    
    def _create_analysis_prompt(self, resume_text: str, job_description: str, job_requirements: Dict) -> str:
        """
        Create a detailed prompt for Gemini AI analysis
        """
        return f"""You are an expert HR recruiter and resume analyst. Analyze the following resume against the job requirements and provide detailed scoring.

JOB DESCRIPTION:
{job_description}

JOB REQUIREMENTS:
- Title: {job_requirements.get('title', 'Not specified')}
- Department: {job_requirements.get('department', 'Not specified')}
- Level: {job_requirements.get('level', 'Not specified')}
- Required Skills: {job_requirements.get('required_skills', [])}
- Experience Level: {job_requirements.get('experience_years', 'Not specified')}
- Education: {job_requirements.get('education', 'Not specified')}

CANDIDATE RESUME:
{resume_text}

Please analyze this resume and provide scoring in the following JSON format only:

{{
    "overall_score": <0-100>,
    "skill_match_score": <0-100>,
    "experience_score": <0-100>,
    "education_score": <0-100>,
    "criteria_analysis": {{
        "technical_skills": {{
            "score": <0-100>,
            "found_skills": ["skill1", "skill2"],
            "missing_skills": ["skill3", "skill4"],
            "reasoning": "explanation"
        }},
        "experience": {{
            "score": <0-100>,
            "years_found": <number>,
            "relevant_experience": ["exp1", "exp2"],
            "reasoning": "explanation"
        }},
        "education": {{
            "score": <0-100>,
            "qualifications": ["qual1", "qual2"],
            "reasoning": "explanation"
        }},
        "soft_skills": {{
            "score": <0-100>,
            "identified_skills": ["skill1", "skill2"],
            "reasoning": "explanation"
        }}
    }},
    "strengths": ["strength1", "strength2", "strength3"],
    "weaknesses": ["weakness1", "weakness2"],
    "recommendation": "HIGHLY_RECOMMENDED|RECOMMENDED|CONSIDER|NOT_RECOMMENDED",
    "detailed_feedback": "Brief explanation of the overall assessment"
}}

Focus on technical skill alignment, relevant work experience, educational background, soft skills, and career progression. Provide only the JSON response without any additional text."""
    
    def _parse_ai_response(self, response_text: str) -> Dict[str, Any]:
        """
        Parse Gemini AI response and extract structured data
        """
        try:
            # Clean the response - remove markdown formatting if present
            cleaned_text = response_text.strip()
            if cleaned_text.startswith('```json'):
                cleaned_text = cleaned_text[7:]
            if cleaned_text.endswith('```'):
                cleaned_text = cleaned_text[:-3]
            
            # Try to extract JSON from the response
            json_match = re.search(r'\{.*\}', cleaned_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                return json.loads(json_str)
            else:
                return self._create_fallback_analysis()
        except Exception as e:
            print(f"Error parsing Gemini response: {e}")
            return self._create_fallback_analysis()
    
    def _create_fallback_analysis(self) -> Dict[str, Any]:
        """
        Create a fallback analysis if AI fails
        """
        return {
            "overall_score": 50,
            "skill_match_score": 50,
            "experience_score": 50,
            "education_score": 50,
            "criteria_analysis": {
                "technical_skills": {
                    "score": 50,
                    "found_skills": [],
                    "missing_skills": [],
                    "reasoning": "Unable to analyze - using fallback scoring"
                },
                "experience": {
                    "score": 50,
                    "years_found": 0,
                    "relevant_experience": [],
                    "reasoning": "Unable to analyze - using fallback scoring"
                },
                "education": {
                    "score": 50,
                    "qualifications": [],
                    "reasoning": "Unable to analyze - using fallback scoring"
                },
                "soft_skills": {
                    "score": 50,
                    "identified_skills": [],
                    "reasoning": "Unable to analyze - using fallback scoring"
                }
            },
            "strengths": ["Resume submitted on time"],
            "weaknesses": ["Unable to analyze resume content"],
            "recommendation": "CONSIDER",
            "detailed_feedback": "Analysis unavailable - manual review recommended"
        }

    def shortlist_candidates(self, job_data: Dict, candidates: List[Dict]) -> List[CandidateScore]:
        """
        Analyze all candidates for a job and return ranked shortlist
        """
        scored_candidates = []
        
        for candidate in candidates:
            try:
                # Get resume text
                resume_text = self._extract_resume_text(candidate)
                
                # Prepare job requirements
                job_requirements = {
                    'title': job_data.get('title', ''),
                    'department': job_data.get('department', ''),
                    'level': job_data.get('level', ''),
                    'required_skills': self._extract_required_skills(job_data),
                    'experience_years': self._extract_experience_requirement(job_data),
                    'education': job_data.get('education_requirement', '')
                }
                
                # Analyze resume
                analysis = self.analyze_resume_for_job(
                    resume_text, 
                    job_data.get('description', ''), 
                    job_requirements
                )
                
                # Create candidate score object
                candidate_score = CandidateScore(
                    candidate_id=str(candidate['_id']),
                    candidate_name=candidate.get('applicantName', 'Unknown'),
                    overall_score=analysis['overall_score'],
                    skill_match_score=analysis['skill_match_score'],
                    experience_score=analysis['experience_score'],
                    education_score=analysis['education_score'],
                    criteria_analysis=analysis['criteria_analysis'],
                    strengths=analysis['strengths'],
                    weaknesses=analysis['weaknesses'],
                    recommendation=analysis['detailed_feedback']                )
                
                scored_candidates.append(candidate_score)
                
            except Exception as e:
                print(f"Error processing candidate {candidate.get('_id', 'unknown')}: {e}")
                continue
        
        # Sort by overall score (highest first)
        scored_candidates.sort(key=lambda x: x.overall_score, reverse=True)
        
        return scored_candidates
    
    def _extract_resume_text(self, candidate: Dict) -> str:
        """
        Extract text from candidate's resume - now with actual file parsing
        """
        resume_parts = []
        
        # Add basic candidate info
        if candidate.get('applicantName'):
            resume_parts.append(f"Name: {candidate['applicantName']}")
        
        if candidate.get('applicantEmail'):
            resume_parts.append(f"Email: {candidate['applicantEmail']}")
        
        # Try to extract text from actual resume file
        resume_file_text = ""
        if candidate.get('resumePath'):
            resume_file_text = self._extract_text_from_file(candidate['resumePath'])
        elif candidate.get('resume_filename'):
            # Try to construct path from filename
            resume_path = f"/uploads/application-resumes/{candidate['resume_filename']}"
            resume_file_text = self._extract_text_from_file(resume_path)
        
        if resume_file_text:
            resume_parts.append("=== RESUME CONTENT ===")
            resume_parts.append(resume_file_text)
            resume_parts.append("=== END RESUME CONTENT ===")
        else:
            # Fallback to database fields if file reading fails
            if candidate.get('experience'):
                resume_parts.append(f"Experience: {candidate['experience']}")
            
            if candidate.get('skills'):
                if isinstance(candidate['skills'], list):
                    resume_parts.append(f"Skills: {', '.join(candidate['skills'])}")
                else:
                    resume_parts.append(f"Skills: {candidate['skills']}")
            
            if candidate.get('education'):
                resume_parts.append(f"Education: {candidate['education']}")
            
            if candidate.get('coverLetter'):
                resume_parts.append(f"Cover Letter: {candidate['coverLetter']}")
        
        final_text = "\n".join(resume_parts) if resume_parts else "No resume information available"
        
        # Debug: Print what we extracted
        print(f"📄 Resume text extracted for {candidate.get('applicantName', 'Unknown')}:")
        print(f"   - File text length: {len(resume_file_text)} characters")
        print(f"   - Total text length: {len(final_text)} characters")
        print(f"   - Resume path: {candidate.get('resumePath', 'N/A')}")
        
        return final_text
    
    def _extract_text_from_file(self, file_path: str) -> str:
        """
        Extract text from resume file (PDF, DOCX, etc.)
        """
        try:
            # Convert relative path to absolute path
            if file_path.startswith('/'):
                file_path = file_path[1:]  # Remove leading slash
            
            # Construct full path
            full_path = os.path.join(os.getcwd(), 'public', file_path.replace('/', os.sep))
            
            print(f"🔍 Trying to read resume file: {full_path}")
            
            if not os.path.exists(full_path):
                print(f"❌ Resume file not found: {full_path}")
                return ""
            
            # Determine file type and extract text
            file_extension = os.path.splitext(full_path)[1].lower()
            
            if file_extension == '.pdf':
                return DocumentProcessor.extract_text_from_pdf(full_path)
            elif file_extension in ['.docx', '.doc']:
                return DocumentProcessor.extract_text_from_docx(full_path)
            elif file_extension == '.txt':
                with open(full_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                print(f"⚠️ Unsupported file type: {file_extension}")
                return ""
                
        except Exception as e:
            print(f"❌ Error extracting text from file {file_path}: {str(e)}")
            return ""
    
    def _extract_required_skills(self, job_data: Dict) -> List[str]:
        """
        Extract required skills from job description
        """
        description = job_data.get('description', '') + " " + job_data.get('jobDescriptionText', '')
        
        # Common tech skills to look for
        common_skills = [
            'JavaScript', 'Python', 'Java', 'React', 'Node.js', 'HTML', 'CSS',
            'SQL', 'MongoDB', 'PostgreSQL', 'AWS', 'Docker', 'Git', 'TypeScript',
            'Angular', 'Vue.js', 'PHP', 'C++', 'C#', 'Ruby', 'Go', 'Kotlin',
            'Swift', 'Flutter', 'React Native', 'Django', 'Flask', 'Express',
            'Spring', 'Laravel', 'Pandas', 'NumPy', 'TensorFlow', 'PyTorch'
        ]
        
        found_skills = []
        description_lower = description.lower()
        
        for skill in common_skills:
            if skill.lower() in description_lower:
                found_skills.append(skill)
        
        return found_skills
    
    def _extract_experience_requirement(self, job_data: Dict) -> str:
        """
        Extract experience requirements from job description
        """
        description = job_data.get('description', '') + " " + job_data.get('jobDescriptionText', '')
        level = job_data.get('level', '')
        
        # Map levels to experience
        level_mapping = {
            'entry': '0-2 years',
            'junior': '1-3 years', 
            'mid': '3-5 years',
            'senior': '5+ years',
            'lead': '7+ years',
            'principal': '10+ years'
        }
        
        return level_mapping.get(level.lower(), level)

# Document processing utilities
class DocumentProcessor:
    """Utility class for processing different document types"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            print(f"Error reading PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading PDF bytes: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(docx_bytes))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error reading DOCX bytes: {str(e)}")
            return ""

# Example usage and testing
if __name__ == "__main__":
    analyzer = AIResumeAnalyzer()
    
    # Test with sample data
    sample_job = {
        'title': 'Senior Frontend Developer',
        'department': 'Engineering',
        'level': 'senior',
        'description': 'We are looking for a Senior Frontend Developer with expertise in React, TypeScript, and modern web technologies.'
    }
    
    sample_candidates = [
        {
            '_id': '12345',
            'applicantName': 'John Doe',
            'applicantEmail': 'john@example.com',
            'skills': ['React', 'JavaScript', 'TypeScript', 'CSS'],
            'experience': '6 years in frontend development',
            'education': 'Bachelor in Computer Science'
        }
    ]
    
    try:
        shortlisted = analyzer.shortlist_candidates(sample_job, sample_candidates)
        print("✅ Gemini AI Analysis completed successfully!")
        for candidate in shortlisted:
            print(f"Candidate: {candidate.candidate_name}, Score: {candidate.overall_score}")
    except Exception as e:
        print(f"❌ Error testing Gemini AI analyzer: {e}")
